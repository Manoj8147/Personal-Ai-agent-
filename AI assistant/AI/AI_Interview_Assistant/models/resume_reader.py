"""
Resume Reader Module
Extracts and structures resume data from PDF, DOCX, and TXT files.
"""

import pdfplumber
from docx import Document
import json
import re
from typing import Dict, List, Any
from pathlib import Path


class ResumeReader:
    """Extract and parse resume information from various file formats."""
    
    # Keywords to identify different sections
    SECTION_KEYWORDS = {
        'education': ['education', 'academic', 'bachelor', 'master', 'degree', 'b.tech', 'b.a', 'mba'],
        'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'proficiency'],
        'experience': ['experience', 'work experience', 'professional experience', 'employment'],
        'projects': ['projects', 'portfolio', 'accomplishments', 'key projects'],
        'certifications': ['certifications', 'certificates', 'credentials'],
        'contact': ['email', 'phone', 'linkedin', 'website', 'address']
    }
    
    def __init__(self, file_path: str):
        """Initialize with file path."""
        self.file_path = file_path
        self.file_extension = Path(file_path).suffix.lower()
        self.raw_text = ""
        
    def extract_text(self) -> str:
        """Extract raw text from file based on extension."""
        try:
            if self.file_extension == '.pdf':
                return self._extract_from_pdf()
            elif self.file_extension == '.docx':
                return self._extract_from_docx()
            elif self.file_extension == '.txt':
                return self._extract_from_txt()
            else:
                raise ValueError(f"Unsupported file type: {self.file_extension}")
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")
    
    def _extract_from_pdf(self) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        return text
    
    def _extract_from_docx(self) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = Document(self.file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        return text
    
    def _extract_from_txt(self) -> str:
        """Extract text from TXT file."""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def parse_resume(self) -> Dict[str, Any]:
        """Parse resume and extract structured information."""
        self.raw_text = self.extract_text()
        
        profile = {
            'name': self._extract_name(),
            'email': self._extract_email(),
            'phone': self._extract_phone(),
            'education': self._extract_education(),
            'skills': self._extract_skills(),
            'experience': self._extract_experience(),
            'projects': self._extract_projects(),
            'certifications': self._extract_certifications(),
            'raw_text': self.raw_text  # Keep raw text for fallback
        }
        
        return profile
    
    def _extract_name(self) -> str:
        """Extract name (usually first line or in header)."""
        lines = self.raw_text.split('\n')
        # Check first few non-empty lines
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) < 50 and not any(
                keyword in line.lower() 
                for keyword in ['email', 'phone', '@', '|', '-']
            ):
                return line
        return "Candidate"
    
    def _extract_email(self) -> str:
        """Extract email address."""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        matches = re.findall(email_pattern, self.raw_text)
        return matches[0] if matches else ""
    
    def _extract_phone(self) -> str:
        """Extract phone number."""
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,9}'
        matches = re.findall(phone_pattern, self.raw_text)
        return matches[0] if matches else ""
    
    def _extract_education(self) -> List[str]:
        """Extract education information."""
        education = []
        lines = self.raw_text.split('\n')
        
        in_education = False
        edu_buffer = []
        
        for line in lines:
            line = line.strip()
            
            # Check if entering education section
            if any(keyword in line.lower() for keyword in self.SECTION_KEYWORDS['education']):
                in_education = True
                continue
            
            # Check if entering other section
            if in_education and line and any(
                keyword in line.lower() 
                for keywords_list in list(self.SECTION_KEYWORDS.values())[1:] 
                for keyword in keywords_list
            ):
                in_education = False
                if edu_buffer:
                    education.extend(edu_buffer)
                    edu_buffer = []
            
            # Collect education lines
            if in_education and line and len(line) < 200:
                # Skip very short lines and special characters only
                if len(line) > 5 and not all(c in '|-/*+=' for c in line):
                    edu_buffer.append(line)
        
        if edu_buffer:
            education.extend(edu_buffer)
        
        return education[:5]  # Limit to 5 entries
    
    def _extract_skills(self) -> List[str]:
        """Extract skills (parse comma/bullet separated skills)."""
        skills = []
        lines = self.raw_text.split('\n')
        
        in_skills = False
        
        for line in lines:
            line = line.strip()
            
            # Check if entering skills section
            if any(keyword in line.lower() for keyword in self.SECTION_KEYWORDS['skills']):
                in_skills = True
                continue
            
            # Check if entering other section
            if in_skills and line and any(
                keyword in line.lower() 
                for keywords_list in list(self.SECTION_KEYWORDS.values())[1:] 
                for keyword in keywords_list
            ):
                in_skills = False
            
            # Extract skills from line
            if in_skills and line and len(line) < 300:
                # Split by comma, semicolon, or bullet points
                skill_items = re.split(r'[,;•\n]', line)
                for skill in skill_items:
                    skill = skill.strip()
                    if skill and len(skill) < 50 and len(skill) > 1:
                        # Remove common prefixes/suffixes
                        skill = re.sub(r'^[-•*]\s*', '', skill)
                        if skill and skill not in skills:
                            skills.append(skill)
        
        return skills[:20]  # Limit to 20 skills
    
    def _extract_experience(self) -> List[Dict[str, str]]:
        """Extract work experience."""
        experience = []
        lines = self.raw_text.split('\n')
        
        in_experience = False
        job_buffer = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Check if entering experience section
            if any(keyword in line.lower() for keyword in self.SECTION_KEYWORDS['experience']):
                in_experience = True
                continue
            
            # Check if entering other section
            if in_experience and line and any(
                keyword in line.lower() 
                for keywords_list in list(self.SECTION_KEYWORDS.values())[1:] 
                for keyword in keywords_list
            ):
                in_experience = False
            
            # Collect experience lines
            if in_experience and line and len(line) < 200:
                job_buffer.append(line)
        
        # Group consecutive lines as one job entry
        current_job = []
        for line in job_buffer:
            if line and (len(current_job) == 0 or len(line) > 10):
                current_job.append(line)
                if len(line) < 50 or any(str(year) in line for year in range(2010, 2030)):
                    # Likely a job title/company line
                    if len(current_job) > 1:
                        experience.append({
                            'title': current_job[0],
                            'description': ' '.join(current_job[1:])
                        })
                        current_job = []
        
        return experience[:5]  # Limit to 5 positions
    
    def _extract_projects(self) -> List[str]:
        """Extract projects."""
        projects = []
        lines = self.raw_text.split('\n')
        
        in_projects = False
        
        for line in lines:
            line = line.strip()
            
            # Check if entering projects section
            if any(keyword in line.lower() for keyword in self.SECTION_KEYWORDS['projects']):
                in_projects = True
                continue
            
            # Check if entering other section
            if in_projects and line and any(
                keyword in line.lower() 
                for keywords_list in list(self.SECTION_KEYWORDS.values())[1:] 
                for keyword in keywords_list
            ):
                in_projects = False
            
            # Extract project
            if in_projects and line and len(line) < 200:
                line = re.sub(r'^[-•*]\s*', '', line)
                if line and line not in projects:
                    projects.append(line)
        
        return projects[:10]  # Limit to 10 projects
    
    def _extract_certifications(self) -> List[str]:
        """Extract certifications."""
        certifications = []
        lines = self.raw_text.split('\n')
        
        in_certs = False
        
        for line in lines:
            line = line.strip()
            
            # Check if entering certifications section
            if any(keyword in line.lower() for keyword in self.SECTION_KEYWORDS['certifications']):
                in_certs = True
                continue
            
            # Check if entering other section
            if in_certs and line and any(
                keyword in line.lower() 
                for keywords_list in list(self.SECTION_KEYWORDS.values())[1:] 
                for keyword in keywords_list
            ):
                in_certs = False
            
            # Extract certification
            if in_certs and line and len(line) < 150:
                line = re.sub(r'^[-•*]\s*', '', line)
                if line and line not in certifications:
                    certifications.append(line)
        
        return certifications[:8]  # Limit to 8 certifications
    
    def to_json(self, profile: Dict[str, Any]) -> str:
        """Convert profile to JSON string."""
        # Remove raw_text for JSON output
        profile_copy = {k: v for k, v in profile.items() if k != 'raw_text'}
        return json.dumps(profile_copy, indent=2)
    
    def to_prompt_context(self, profile: Dict[str, Any]) -> str:
        """Convert profile to formatted text for AI prompts."""
        context = f"""CANDIDATE PROFILE
================

Name: {profile.get('name', 'Candidate')}
Email: {profile.get('email', 'Not provided')}
Phone: {profile.get('phone', 'Not provided')}

EDUCATION:
{self._format_list(profile.get('education', []))}

SKILLS:
{self._format_list(profile.get('skills', []))}

EXPERIENCE:
{self._format_experience(profile.get('experience', []))}

PROJECTS:
{self._format_list(profile.get('projects', []))}

CERTIFICATIONS:
{self._format_list(profile.get('certifications', []))}
"""
        return context
    
    @staticmethod
    def _format_list(items: List[Any], prefix: str = "• ") -> str:
        """Format list of items with bullet points."""
        if not items:
            return "None provided"
        return '\n'.join(f"{prefix}{item}" for item in items)
    
    @staticmethod
    def _format_experience(experience: List[Dict[str, str]]) -> str:
        """Format experience entries."""
        if not experience:
            return "None provided"
        formatted = []
        for exp in experience:
            formatted.append(f"• {exp.get('title', 'Position')}")
            if exp.get('description'):
                formatted.append(f"  {exp.get('description')}")
        return '\n'.join(formatted)


def parse_resume_file(file_path: str) -> Dict[str, Any]:
    """Convenience function to parse a resume file."""
    reader = ResumeReader(file_path)
    return reader.parse_resume()
